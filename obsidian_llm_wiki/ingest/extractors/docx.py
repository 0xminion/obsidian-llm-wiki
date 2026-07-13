"""DOCX extractor — extracts text from Word .docx files using python-docx.

Dependency (optional): ``python-docx``.
Install with: ``pip install obsidian-llm-wiki[docx]``
"""

from __future__ import annotations

import logging
from pathlib import Path

from obsidian_llm_wiki.core.models import SourceDoc
from obsidian_llm_wiki.ingest.extractors import register_extractor

logger = logging.getLogger("obswiki.ingest.extractors.docx")

# ── Dependency check ────────────────────────────────────────────────────

try:
    import docx  # python-docx  # type: ignore[import-untyped]

    _DEPS_AVAILABLE = True
except ImportError:
    _DEPS_AVAILABLE = False


# ── Matching ────────────────────────────────────────────────────────────


def _is_docx(parsed, raw: str) -> bool:
    """Match .docx file paths."""
    return raw.lower().endswith(".docx")


# ── Registration (only if deps available) ────────────────────────────────

if _DEPS_AVAILABLE:

    @register_extractor(_is_docx)
    def extract_docx(raw_path: str) -> SourceDoc:
        """Extract text from a Word .docx file.

        Iterates paragraphs and tables, concatenating text content.
        Title is derived from core properties or the first non-empty paragraph.
        """
        path = Path(raw_path)
        if not path.is_file():
            raise RuntimeError(f"DOCX file not found: {raw_path}")

        document = docx.Document(str(path))  # type: ignore[union-attr]

        # ── Title from core properties ──────────────────────────────────
        title = ""
        core_props = document.core_properties
        if core_props and core_props.title:
            title = core_props.title.strip()
        if not title:
            # Use first non-empty paragraph.
            for para in document.paragraphs:
                if para.text.strip():
                    title = para.text.strip()[:120]
                    break
        if not title:
            title = path.stem

        # ── Extract text from paragraphs ─────────────────────────────────
        parts: list[str] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                # Prefix headings with markdown # for structure.
                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading 1" in style_name:
                    parts.append(f"# {text}")
                elif "heading 2" in style_name:
                    parts.append(f"## {text}")
                elif "heading 3" in style_name:
                    parts.append(f"### {text}")
                elif "heading" in style_name:
                    parts.append(f"#### {text}")
                else:
                    parts.append(text)

        # ── Extract text from tables ────────────────────────────────────
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        content = "\n\n".join(parts)
        if not content.strip():
            raise RuntimeError(f"DOCX has no extractable text: {raw_path}")

        return SourceDoc(
            title=title,
            content=content,
            url=str(path),
        )
